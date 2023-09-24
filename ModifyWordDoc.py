from docx import Document
from docx.shared import RGBColor

def modify_word_document(file_path):
    doc = Document(file_path)
    
    # Modify content
    for paragraph in doc.paragraphs:
        if 'replace' in paragraph.text:
            paragraph.text = paragraph.text.replace('replace', 'replacement')
    
    # Modify formatting
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.italic = True
            run.underline = True
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # Red color
    
    # Save the modified document
    doc.save(file_path)

# Usage example
existing_document_path = 'path/to/existing_document.docx'
modify_word_document(existing_document_path)
