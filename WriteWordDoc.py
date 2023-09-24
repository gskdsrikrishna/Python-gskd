from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

def write_word_document(file_path):
    doc = Document()
    
    # Add text and formatting
    doc.add_paragraph('Hello, World!').bold = True
    paragraph = doc.add_paragraph('This is a new paragraph.')
    paragraph.add_run('This text is italic.').italic = True
    paragraph.add_run(' This text is underlined.').underline = True
    
    # Save the document
    doc.save(file_path)

# Usage example
new_document_path = 'path/to/new_document.docx'
write_word_document(new_document_path)
