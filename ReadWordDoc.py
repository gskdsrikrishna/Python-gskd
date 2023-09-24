from docx import Document

def read_word_document(file_path):
    doc = Document(file_path)
    
    # Extract text
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    
    # Extract formatting information
    formatting = []
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            formatting.append((run.text, run.bold, run.italic, run.underline))
    
    return text, formatting

# Usage example
document_path = 'path/to/document.docx'
text, formatting = read_word_document(document_path)
print(text)
print(formatting)
