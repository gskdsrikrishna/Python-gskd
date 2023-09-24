from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def add_watermark_to_word(input_file, output_file, watermark_text=None, watermark_image=None):
    doc = Document(input_file)
    watermark = None
    
    if watermark_text:
        watermark = doc.sections[0].footer.paragraphs[0]
        watermark.text = watermark_text
        watermark.alignment = 1  # Center alignment
    
    if watermark_image:
        if watermark is None:
            watermark = doc.sections[0].footer.paragraphs[0]
        
        run = watermark.add_run()
        run.add_picture(watermark_image)
        for paragraph in watermark.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(48)  # Adjust the font size as needed
                run.font.color.theme_color = MSO_THEME_COLOR.LIGHT_2  # Adjust the font color as needed
    
    for section in doc.sections:
        section._sectPr.xpath('./w:headerReference')[0].attrib.clear()
        section._sectPr.xpath('./w:footerReference')[0].attrib.clear()
        section._sectPr.xpath('./w:headerReference')[0].append(parse_xml('<w:headerReference w:type="default" />'))
    
    doc.save(output_file)

# Usage example
input_path = 'path/to/document.docx'
output_path = 'path/to/watermarked_document.docx'
watermark_text = 'Confidential'
watermark_image = 'path/to/watermark_image.png'
add_watermark_to_word(input_path, output_path, watermark_text, watermark_image)
