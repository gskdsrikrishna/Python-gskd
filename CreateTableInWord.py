from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def create_table_in_word(output_file):
    doc = Document()
    
    # Create a table
    table = doc.add_table(rows=3, cols=3)
    
    # Set table style
    table.style = 'Table Grid'
    
    # Add data to the table
    table.cell(0, 0).text = 'Header 1'
    table.cell(0, 1).text = 'Header 2'
    table.cell(0, 2).text = 'Header 3'
    
    table.cell(1, 0).text = 'Cell 1'
    table.cell(1, 1).text = 'Cell 2'
    table.cell(1, 2).text = 'Cell 3'
    
    table.cell(2, 0).text = 'Cell 4'
    table.cell(2, 1).text = 'Cell 5'
    table.cell(2, 2).text = 'Cell 6'
    
    # Customize table formatting
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(100)  # Adjust the cell width as needed
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    run.bold = True  # Apply bold formatting to the text
    
    doc.save(output_file)

# Usage example
output_path = 'path/to/table_document.docx'
create_table_in_word(output_path)
